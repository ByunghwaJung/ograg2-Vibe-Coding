from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- 기본 스타일 설정 ---
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    set_font(run, bold=True, size=16 if level == 1 else (13 if level == 2 else 11))
    if level == 1:
        run.font.color.rgb = RGBColor(0x5B, 0x2D, 0x8E)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, bold=bold, size=10)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = '맑은 고딕'
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            cells[c].text = val
            cells[c].paragraphs[0].runs[0].font.name = '맑은 고딕'
            cells[c].paragraphs[0].runs[0].font.size = Pt(10)

# =====================================================================
# 표지
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('OG-RAG2 구현 보고서')
set_font(r, bold=True, size=24, color=(0x5B, 0x2D, 0x8E))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Ontology-Grounded Retrieval-Augmented Generation\nWindows 환경 실습 및 분석')
set_font(r2, size=13, color=(0x44, 0x44, 0x44))

doc.add_paragraph()
doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info_p.add_run('성균관대학교 스마트팩토리융합학과\n2026년 5월')
set_font(r3, size=11)

doc.add_page_break()

# =====================================================================
# 1. 개요
# =====================================================================
heading('1. 개요')
body('OG-RAG(Ontology-Grounded Retrieval-Augmented Generation)는 도메인 특화 온톨로지를 활용하여 LLM의 응답 정확도와 사실 기반성을 향상시키는 시스템이다. 본 보고서는 Microsoft 공개 소스인 ograg2를 Windows 11 환경에서 실제로 설치·실행하고, 발생한 오류를 수정하여 동작시킨 전 과정을 기술한다.')

doc.add_paragraph()
body('핵심 특징:', bold=True)
bullet('온톨로지 기반 구조적 지식 검색 (Ontology-Grounded Retrieval)')
bullet('Hypergraph를 활용한 복합 관계 표현 (Hyperedge 구조)')
bullet('Greedy Set Cover 알고리즘 기반 최적 컨텍스트 검색 (retriever.py)')
bullet('농업 도메인(대두, 밀) 특화 지식 그래프 구축 및 질의응답')

doc.add_page_break()

# =====================================================================
# 2. 실행 환경 구성
# =====================================================================
heading('2. 실행 환경 구성')

heading('2.1 환경 스펙', level=2)
add_table(
    ['항목', '내용'],
    [
        ['OS', 'Windows 11 Enterprise'],
        ['Python', '3.11.9'],
        ['IDE', 'Visual Studio Code'],
        ['LLM', 'GPT-4o-mini (OpenAI API)'],
        ['Embedding', 'text-embedding-3-small (OpenAI API)'],
        ['가상환경', 'venv (Python 내장)'],
    ]
)

doc.add_paragraph()
heading('2.2 설치 절차', level=2)
body('① 가상환경 생성 및 활성화')
code_block('python -m venv venv\n.\\venv\\Scripts\\activate')
body('② 패키지 설치')
code_block('pip install -r requirements.txt')
body('③ API 키 설정 (api_keys.yaml)')
code_block('OPENAI_API_KEY: sk-proj-...')

heading('2.3 Windows 환경 의존성 충돌 해결', level=2)
body('requirements.txt는 Linux 환경을 기준으로 작성되어 Windows에서 다수의 의존성 충돌이 발생하였다. 주요 해결 방법은 다음과 같다.')
bullet('nvidia-cuda 관련 패키지: Windows CPU 환경에서는 설치 스킵 또는 cpu 버전으로 대체')
bullet('SecretStorage, jeepney: Linux 전용 패키지로 Windows에서 무시')
bullet('직접 pip install로 충돌 패키지 버전 조정 후 실행 성공')

doc.add_page_break()

# =====================================================================
# 3. 소스코드 구조 분석
# =====================================================================
heading('3. 소스코드 구조 분석')

heading('3.1 디렉터리 구조', level=2)
code_block(
    'ograg2/\n'
    '├── ontology_mapping/      - 온톨로지 로드 및 개념 계층 파싱\n'
    '├── knowledge_graph/       - KG 트리플 생성 (Factual Block → Hyperedge)\n'
    '├── query_engine/          - Query 처리 및 응답 생성\n'
    '│   ├── ontograph_query_engine.py  - Hypergraph 구성 및 관리\n'
    '│   ├── retriever.py               - Greedy Set Cover 검색\n'
    '│   └── knowledge_graph_query_engine.py\n'
    '├── utils/                 - 공통 유틸리티 (LLM 로드, 문서 파싱)\n'
    '├── data/\n'
    '│   ├── md/sample/         - 입력 문서 (Markdown)\n'
    '│   ├── kg/sample/         - 생성된 KG (.pkl, .npy)\n'
    '│   └── ontology/          - 온톨로지 JSON-LD\n'
    '├── config.yaml            - 실행 설정\n'
    '├── build_knowledge_graph.py\n'
    '└── query_llm.py'
)

heading('3.2 핵심 모듈 역할', level=2)
add_table(
    ['모듈', '역할'],
    [
        ['ontology_mapping.py', '온톨로지 JSON-LD 로드 및 도메인 개념 계층 파싱, LLM 호출로 노드 생성'],
        ['knowledge_graph.py', '온톨로지 노드 기반 KG 트리플 생성, Factual Block → Hyperedge 변환'],
        ['ontograph_query_engine.py', 'Hypergraph 구성 및 Hyperedge 관리, 쿼리 처리 메인 엔진'],
        ['retriever.py', 'Greedy Set Cover 기반 최적 컨텍스트 검색'],
        ['query_llm.py', 'NL 쿼리 수신 → 쿼리 엔진 호출 → LLM 응답 생성 출력'],
    ]
)

heading('3.3 실행 파이프라인', level=2)
body('OG-RAG의 전체 실행 흐름은 5단계로 구성된다.')
add_table(
    ['단계', '처리 내용', '관련 모듈'],
    [
        ['① 문서 로드', 'data/md/sample/*.md → DocumentLoader', 'utils/utils.py'],
        ['② 온톨로지', 'LLM 호출 → 개념 계층 구조 추출', 'ontology_mapping.py'],
        ['③ Factual Block', 'Algorithm 1: 문장 → 의미 단위 분리', 'knowledge_graph.py'],
        ['④ Hypergraph', 'Block 간 관계 → Hyperedge 생성', 'ontograph_query_engine.py'],
        ['⑤ 응답 생성', 'Greedy Set Cover → LLM 호출 → 최종 답변', 'retriever.py'],
    ]
)

doc.add_page_break()

# =====================================================================
# 4. Knowledge Graph 구축 실습
# =====================================================================
heading('4. Ontology 기반 Knowledge Graph 구축 실습')

heading('4.1 실행 명령어', level=2)
code_block('python build_knowledge_graph.py --config_file config.yaml')

heading('4.2 처리 흐름', level=2)
add_table(
    ['단계', '설명'],
    [
        ['① 문서 로드', 'data/md/sample/*.md 파일을 텍스트로 파싱'],
        ['② 온톨로지 생성', 'LLM 호출로 ontology_node_0.jsonld 생성'],
        ['③ Factual Block', '문장 단위 의미 분리 및 구조화'],
        ['④ Hyperedge 생성', 'Block 간 관계를 Hyperedge로 변환'],
        ['⑤ KG 저장', 'data/kg/sample/ontology_triples.pkl 저장'],
    ]
)

heading('4.3 실행 로그 (요약)', level=2)
code_block(
    'Processing subdirectory: data\\md\\sample\n'
    'Processing Node 0. Generating data for ontology.\n'
    'Saved ontology data from node 0 to data/kg/sample/ontology\\ontology_node_0.jsonld\n'
    'Started processing: ontology\n'
    'Finished processing file data/kg/sample\\ontology\\ontology_node_0.jsonld\n'
    'All triples with document names for ontology have been saved to data/kg/sample\\ontology_triples.pkl'
)

heading('4.4 생성된 KG 파일', level=2)
add_table(
    ['파일명', '크기', '설명'],
    [
        ['ontology_triples.pkl', '5,522 bytes', 'KG 트리플 데이터 (subject, predicate, object)'],
        ['ontology_node_0.jsonld', '5,621 bytes', '온톨로지 노드 JSON-LD'],
    ]
)

heading('4.5 발생 오류 및 수정', level=2)
body('build_knowledge_graph.py 실행 중 두 가지 버그를 발견하고 직접 수정하였다.')
doc.add_paragraph()
body('① UnboundLocalError (knowledge_graph.py:161)', bold=True)
bullet('원인: except 블록에서 json_filename 변수가 미할당 상태로 참조됨')
bullet('수정: futures 리스트 대신 {future: filename} 딕셔너리로 변경하여 항상 파일명 참조 가능하도록 수정')
doc.add_paragraph()
body('② SyntaxError: unterminated string literal (knowledge_graph.py:202)', bold=True)
bullet('원인: LLM이 트리플 문자열 내부에 개행문자를 포함한 응답을 반환')
bullet('수정 1: _clean_triples_str() 메서드 추가 - 마크다운 코드 펜스 제거 및 개행문자 공백 치환')
bullet('수정 2: _extract_triples_regex() 폴백 추가 - ast.literal_eval 실패 시 정규식으로 트리플 추출')

doc.add_page_break()

# =====================================================================
# 5. OG-RAG Query 실행 결과
# =====================================================================
heading('5. OG-RAG Query 실행 결과')

heading('5.1 실행 명령어', level=2)
code_block('python query_llm.py --config_file config.yaml')

heading('5.2 질의응답 결과', level=2)
body('Q1 — 한국어 질의', bold=True)
body('"대두 저장 조건을 설명해줘"', indent=True)
doc.add_paragraph()
body('응답 (Response):', bold=True)
body(
    '대두 저장 조건은 다음과 같습니다:\n'
    '1. 저장 장소는 시원하고 건조해야 하며, 통풍이 잘 되어야 하고 해충이 없어야 합니다.\n'
    '2. 대두 품질은 가능한 한 새로도 보관해야 합니다.\n'
    '3. 적재를 할 경우, 4~5개의 포대로 쌓아야 하며, 높이 최대 1.5m를 넘기지 않도록 해야 합니다.\n'
    '   이는 대두 씨앗의 생존력과 발아율을 유지하기 위함입니다.',
    indent=True
)
doc.add_paragraph()

body('Q2 — 영어 질의', bold=True)
body('"What is soybean rust?"', indent=True)
doc.add_paragraph()
body('응답 (Response):', bold=True)
body(
    'Soybean rust is a disease caused by fungal pathogens that affects soybean plants. '
    'The symptoms of soybean rust include the appearance of chlorotic gray-brown spots on the leaves, '
    'which are more abundant on the lower surface. As the disease progresses, these spots increase in size '
    'and form pustules, leading to the browning of leaves and early defoliation. '
    'This can result in a reduction in the number of pods, seeds, and seed weight. '
    'The presence of loose brown powder from ruptured pustules is a characteristic symptom of soybean rust.',
    indent=True
)

heading('5.3 결과 분석', level=2)
bullet('한국어 및 영어 질의 모두 정상 처리됨을 확인')
bullet('온톨로지 기반 구조화된 지식을 활용하여 도메인 전문 답변 생성')
bullet('대두 저장 조건 질의에서 구체적 수치(1.5m, 4~5포대)까지 포함한 정확한 응답 확인')
bullet('soybean rust 질의에서 증상, 진행 과정, 영향 등 계층적 관계를 포함한 응답 생성')

doc.add_page_break()

# =====================================================================
# 6. Baseline RAG vs OG-RAG 비교
# =====================================================================
heading('6. Baseline RAG vs Ontology Graph RAG 비교')

add_table(
    ['항목', '일반 RAG (Baseline)', 'OG-RAG (Proposed)', '개선 효과'],
    [
        ['검색 방식', 'Vector Similarity', 'Ontology + Hypergraph', '구조적 관계 반영 ↑'],
        ['관계 이해', '제한적 (단편적 체크)', '강화 (계층 관계 활용)', '복잡 질의 처리 ↑'],
        ['Multi-hop', '어려움', '가능 (Hyperedge 순회)', '추론 깊이 ↑'],
        ['설명 가능성', '낮음', '높음 (온톨로지 경로)', '신뢰성 ↑'],
        ['도메인 적합', '범용', '도메인 특화 구조화', 'Context Recall ↑'],
    ]
)

doc.add_paragraph()
heading('6.1 일반 RAG 한계', level=2)
bullet('청크 단위 독립 검색 → 관계 단절')
bullet('개념 상위/하위 관계 무시')
bullet('동의어/변형 표현 처리 약함')
bullet('복잡한 인과 관계 질의 실패')

heading('6.2 OG-RAG 강점', level=2)
bullet('온톨로지 계층 기반 개념 확장 검색')
bullet('Hyperedge로 복잡 관계 커버')
bullet('Greedy Set Cover → 최소 중복 최대 커버')
bullet('도메인 전문 지식 구조화 활용 가능')

doc.add_page_break()

# =====================================================================
# 7. 실행 결과 분석 및 한계
# =====================================================================
heading('7. 실행 결과 분석 및 한계')

heading('7.1 잘된 점', level=2)
bullet('실제 공개 소스(ograg2) 실행 성공')
bullet('Windows 환경 dependency 충돌 직접 해결')
bullet('Ontology 기반 구조화 응답 생성 확인')
bullet('한국어 질의 처리 가능 확인')
bullet('KG 파일 (.pkl/.npy) 생성 정상 완료')

heading('7.2 한계점', level=2)
bullet('Windows dependency 충돌 다수 발생 — Linux/Docker 환경 권장')
bullet('일부 소스 f-string syntax 오류 존재 (Python 버전 의존)')
bullet('Together AI 의존성 → 사용 불가 (API 키 없음)')
bullet('실행 환경 재현 난이도 높음 — 환경 스크립트 부재')
bullet('대용량 도메인 KG 구축 시간 과다 — 배치 처리 최적화 필요')

heading('7.3 개선 방향', level=2)
bullet('Docker / Linux 기반 배포 환경 구성')
bullet('requirements.txt Windows 호환 버전 정리 및 분리')
bullet('로컬 LLM (Ollama) 연동으로 API 비용 절감')
bullet('제조/MES 도메인 온톨로지 적용 검토')
bullet('CI/CD 기반 자동 환경 구성 스크립트화')

doc.add_page_break()

# =====================================================================
# 8. MES 기반 확장 가능성
# =====================================================================
heading('8. MES 기반 Ontology Graph RAG 확장 가능성')

heading('8.1 MES 핵심 데이터 엔티티', level=2)
add_table(
    ['엔티티', '설명'],
    [
        ['CIM_LOT', '생산 Lot 단위 (Lot ID, 상태)'],
        ['CIM_PRODUCTORDER', '생산 지시 (공정, 수량, 일정)'],
        ['CIM_PROCESSSEGMENT', '공정 세그먼트 (스텝, 조건)'],
        ['CIM_MATERIALLOT', '자재 Lot (BOM, 투입 이력)'],
        ['CIM_DEFECT', '불량 레코드 (유형, 위치, 시간)'],
        ['CIM_EQUIPMENT', '설비 정보 (상태, 가동률)'],
    ]
)

heading('8.2 확장 파이프라인', level=2)
add_table(
    ['단계', '내용'],
    [
        ['MES DB (SQL Server)', 'CIM 테이블에서 생산 데이터 추출'],
        ['Manufacturing Ontology', '공정·설비·불량 개념 계층 정의'],
        ['Ontology Hypergraph', '엔티티 간 복합 관계를 Hyperedge로 구성'],
        ['Graph RAG Retriever', 'Greedy Set Cover 기반 관련 컨텍스트 검색'],
        ['NeuroFactory AI Assistant', 'NL 질의 → 제조 지식 기반 응답 생성'],
    ]
)

heading('8.3 기대 효과', level=2)
body('① 자연어 MES 조회', bold=True)
body('"오늘 1라인 불량 현황 알려줘" → 온톨로지 기반 SQL 생성', indent=True)
doc.add_paragraph()
body('② Multi-hop 추론', bold=True)
body('LOT → DEFECT → EQUIPMENT 연결로 근본 원인 분석 가능', indent=True)
doc.add_paragraph()
body('③ 스마트팩토리 확장', bold=True)
body('YG-1 생산 데이터 기반 → 실시간 AI 의사결정 지원', indent=True)

doc.add_page_break()

# =====================================================================
# 9. 순차적 구현 매뉴얼
# =====================================================================
heading('9. 순차적 구현 매뉴얼')
body('아래 순서대로 따라하면 Windows 환경에서 OG-RAG2를 처음부터 실행할 수 있다.')
doc.add_paragraph()

heading('Step 1. 소스코드 준비', level=2)
code_block('git clone https://github.com/microsoft/ograg2.git\ncd ograg2')

heading('Step 2. 가상환경 생성 및 활성화', level=2)
code_block('python -m venv venv\n.\\venv\\Scripts\\activate')
body('※ 터미널 좌측에 (venv) 프롬프트가 표시되면 활성화 성공')

heading('Step 3. 패키지 설치', level=2)
code_block('pip install -r requirements.txt')
body('※ Windows에서 일부 패키지(SecretStorage, nvidia-cuda 등) 설치 실패 시 무시하고 진행')

heading('Step 4. API 키 설정', level=2)
body('api_keys.yaml 파일 생성 후 OpenAI API 키 입력:')
code_block('OPENAI_API_KEY: sk-proj-여기에_발급받은_키_입력')
body('config.yaml에서 모델 및 데이터 경로 설정:')
code_block(
    'model:\n'
    '  api_base: https://api.openai.com/v1\n'
    '  deployment_name: gpt-4o-mini\n'
    '  api_type: openai\n\n'
    'data:\n'
    '  documents_dir: data/md/sample\n'
    '  ontology_path: data/ontology/farm_cropcultivation_schema_ontology_jsonld.json\n'
    '  kg_storage_path: data/kg/sample'
)

heading('Step 5. 입력 문서 준비', level=2)
body('data/md/sample/ 폴더에 분석할 Markdown 문서 배치:')
code_block('# 예시: 기존 soybean 요약 텍스트 활용\nCopy-Item data\\kg\\soybean\\summarized_texts.md data\\md\\sample\\')

heading('Step 6. Knowledge Graph 구축', level=2)
code_block('python build_knowledge_graph.py --config_file config.yaml')
body('정상 실행 시 출력 예시:')
code_block(
    'Processing subdirectory: data\\md\\sample\n'
    'Processing Node 0. Generating data for ontology.\n'
    'Saved ontology data from node 0 to data/kg/sample/ontology\\ontology_node_0.jsonld\n'
    'Finished processing file data/kg/sample\\ontology\\ontology_node_0.jsonld\n'
    'All triples with document names for ontology have been saved to data/kg/sample\\ontology_triples.pkl'
)
body('생성 파일 확인:')
code_block('Get-ChildItem data\\kg\\sample\\ -Recurse | Format-Table Name, Length')

heading('Step 7. 질의응답 실행', level=2)
code_block('python query_llm.py --config_file config.yaml')
body('프롬프트 입력 화면:')
code_block('Type your query (press Enter to stop): 대두 저장 조건을 설명해줘')
body('빈 엔터 입력 시 종료')

heading('Step 8. 버그 수정 (Windows 환경)', level=2)
body('아래 두 버그는 Windows 환경에서 발생하며 직접 수정이 필요하다.')
doc.add_paragraph()
body('① knowledge_graph/knowledge_graph.py — UnboundLocalError 수정', bold=True)
body('수정 전: futures 리스트로 파일명 조회 → except 블록에서 미할당 변수 참조')
body('수정 후: {future: filename} 딕셔너리로 변경')
code_block(
    '# 수정 전\n'
    'futures = [executor.submit(...) for json_filename in json_filenames]\n'
    'for future in as_completed(futures):\n'
    '    try:\n'
    '        ...\n'
    '        json_filename = json_filenames[futures.index(future)]\n'
    '    except Exception as e:\n'
    '        LOGGER.error(f"... {json_filename}: {e}")  # ← 오류 발생\n\n'
    '# 수정 후\n'
    'future_to_filename = {executor.submit(..., f): f for f in json_filenames}\n'
    'for future in as_completed(future_to_filename):\n'
    '    json_filename = future_to_filename[future]  # ← 항상 할당됨\n'
    '    try:\n'
    '        ...\n'
    '    except Exception as e:\n'
    '        LOGGER.error(f"... {json_filename}: {e}")'
)
doc.add_paragraph()
body('② knowledge_graph/knowledge_graph.py — LLM 응답 파싱 오류 수정', bold=True)
body('수정: _clean_triples_str() 및 _extract_triples_regex() 메서드 추가')
code_block(
    'def _clean_triples_str(self, triples_str):\n'
    '    import re\n'
    '    triples_str = re.sub(r\'^```(?:python)?\\s*\', \'\', triples_str.strip())\n'
    '    triples_str = re.sub(r\'\\s*```$\', \'\', triples_str)\n'
    '    return triples_str.replace(\'\\r\\n\',\' \').replace(\'\\n\',\' \').strip()\n\n'
    'def _extract_triples_regex(self, triples_str):\n'
    '    import re\n'
    '    pattern = r"""\\(\\s*[\'"]([^\'"]+)[\'"],\\s*[\'"]([^\'"]+)[\'"],\\s*[\'"]([^\'"]+)[\'"]\\s*\\)"""\n'
    '    return [(s,p,o) for s,p,o in re.findall(pattern, triples_str)]'
)

doc.add_page_break()

# =====================================================================
# 10. LLM 프롬프트 입력 내용
# =====================================================================
heading('10. LLM 프롬프트 입력 내용')
body('OG-RAG2가 내부적으로 LLM에 전달하는 주요 프롬프트 템플릿을 기술한다.')
doc.add_paragraph()

heading('10.1 Knowledge Graph 트리플 생성 프롬프트', level=2)
body('파일 위치: knowledge_graph/knowledge_graph.py (KG_TRIPLET_ONTOLOGY_EXTRACT_TMPL)')
body('역할: 온톨로지 JSON-LD를 입력받아 (subject, predicate, object) 트리플 리스트를 생성')
doc.add_paragraph()
code_block(
    'Using the @graph namespace in the following json-ld, generate a complete python list\n'
    'of tuples of triples for knowledge graph in the format (subject, predicate, object).\n'
    'Keep the property names exactly as it is in the Json-ld.\n'
    'The subject, predicate, and object can only be strings.\n'
    'Subjects and objects should be in natural language.\n'
    'Make sure that the predicate is structured so that it is a grammatically correct phrase.\n'
    'The triples cannot be nested, so please flatten them.\n'
    'For nested structure, flatten by rearranging keys in a grammatically sensible way.\n'
    'Generate all triples.\n'
    'Do not add any other text in response other than the list of tuples of triples.\n'
    '------------------------------\n'
    'JSON-LD:\n'
    '{data}'
)
doc.add_paragraph()
body('입력 예시 ({data} 부분):')
code_block(
    '{\n'
    '  "@graph": [\n'
    '    {\n'
    '      "@type": "Crop",\n'
    '      "name": "Soybean",\n'
    '      "has_storage_conditions": {\n'
    '        "temperature": "cool and dry",\n'
    '        "max_stack_height": "1.5m"\n'
    '      }\n'
    '    }\n'
    '  ]\n'
    '}'
)
body('출력 예시 (LLM 응답):')
code_block(
    "[('Soybean', 'has storage condition temperature', 'cool and dry'),\n"
    " ('Soybean', 'has storage condition max stack height', '1.5m')]"
)

heading('10.2 온톨로지 매핑 프롬프트', level=2)
body('파일 위치: ontology_mapping/ontology_mapping.py')
body('역할: 입력 문서 청크를 온톨로지 개념에 매핑하여 JSON-LD 노드 생성')
code_block(
    'Given the following text and the ontology schema, extract and map the relevant\n'
    'information into the ontology structure as a JSON-LD format.\n\n'
    'Ontology Schema:\n'
    '{ontology_schema}\n\n'
    'Text:\n'
    '{text}\n\n'
    'Output only valid JSON-LD. Do not include any explanation.'
)

heading('10.3 질의응답 프롬프트', level=2)
body('파일 위치: query_engine/ontograph_query_engine.py')
body('역할: 검색된 Hyperedge 컨텍스트를 포함하여 LLM에 최종 답변 요청')
code_block(
    'You are an expert assistant. Use the following context retrieved from a knowledge graph\n'
    'to answer the question accurately.\n\n'
    'Context (Ontology-grounded knowledge):\n'
    '{context}\n\n'
    'Question: {question}\n\n'
    'Answer based only on the provided context. If the context does not contain\n'
    'sufficient information, state that clearly.'
)

doc.add_page_break()

# =====================================================================
# 11. 결론
# =====================================================================
heading('11. 결론')
heading('11. 결론')
body(
    '본 실습에서는 Microsoft의 OG-RAG2 공개 소스를 Windows 11 환경에서 설치·실행하고, '
    '발생한 버그를 직접 수정하여 동작시켰다. 온톨로지 기반 지식 그래프 구축부터 '
    'Hyperedge 기반 검색, LLM 응답 생성까지 전체 파이프라인을 실증하였다.'
)
doc.add_paragraph()
body(
    '일반 RAG와 달리 OG-RAG는 도메인 온톨로지를 활용한 구조적 지식 검색이 가능하여 '
    '농업·제조·의료 등 전문 도메인에서 높은 응답 정확도를 기대할 수 있다. '
    '향후 MES 데이터와 연계한 스마트팩토리 AI 어시스턴트로 확장할 경우 '
    '실질적인 현장 의사결정 지원 시스템 구현이 가능할 것으로 판단된다.'
)

# 저장
output_path = r'c:\AI\ograg2\OG-RAG2_구현보고서.docx'
doc.save(output_path)
print(f"저장 완료: {output_path}")
